import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import io
import datetime
import time
import json
from google import genai
from google.genai import types

# --- Helper Functions for Word Styling ---
def set_cell_background(cell, fill_hex):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def format_cell_text(cell, font_size=9.5, bold=False, alignment=None):
    for p in cell.paragraphs:
        if alignment:
            p.alignment = alignment
        for r in p.runs:
            r.font.size = Pt(font_size)
            r.font.name = 'Arial'
            r.font.bold = bold


# --- Resilient Evaluation Content Generation Engine ---
def generate_lesson_evaluation_with_gemini(client, topic, syllabus_code):
    if not topic.strip():
        return "", "", ""

    safe_topic = topic.replace('"', "'").strip()

    fallback_www = f"Students encountered foundational layout and core conceptual clarity challenges when mapping out '{safe_topic}'."
    fallback_ebw = f"Utilize target review exemplars, scaffolded structural tracking tasks, or collaborative breakdowns of '{safe_topic}' rules."
    fallback_wf = f"Reinforce core organizational elements and logical structures of '{safe_topic}' in the opening task. Carried out next lesson."

    max_retries = 3
    for attempt in range(max_retries):
        try:
            prompt = f"""
            Analyze the following lesson topic taught under syllabus specification "{syllabus_code or 'General Curriculum'}".
            Topic: "{safe_topic}"

            Provide a custom, analytical pedagogical evaluation tailored specifically to this topic. Return a JSON object with these exact keys:
            - "www": A unique, context-specific student struggle, misconception, or mistake encountered when learning this exact topic (approx 25 words).
            - "ebw": A creative, hands-on classroom task, assignment, or review activity to directly target this issue (approx 25 words).
            - "wf": A remediation path. MUST end strictly with the phrase: Carried out next lesson. (approx 20 words).
            """

            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema={
                        "type": "OBJECT",
                        "properties": {
                            "www": {"type": "STRING"},
                            "ebw": {"type": "STRING"},
                            "wf": {"type": "STRING"}
                        },
                        "required": ["www", "ebw", "wf"]
                    },
                    temperature=0.7
                )
            )

            raw_text = response.text.strip()
            if raw_text.startswith("```"):
                raw_text = raw_text.split("```")[1]
                if raw_text.startswith("json"):
                    raw_text = raw_text[4:]

            data = json.loads(raw_text.strip())
            www = data.get("www", "").strip()
            ebw = data.get("ebw", "").strip()
            wf = data.get("wf", "").strip()

            if wf and "Carried out next lesson." not in wf:
                wf = wf.rstrip('.') + ". Carried out next lesson."

            if www and ebw and wf:
                return www, ebw, wf

        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(2.5)
                continue
            break

    return fallback_www, fallback_ebw, fallback_wf


def generate_weekly_conclusion_with_gemini(client, topics, syllabus_code):
    valid_topics = [t for t in topics if t.strip()]
    if not valid_topics:
        return "No lessons were conducted during this evaluation timeframe."
    try:
        prompt = f"Write a professional, concise summary conclusion (around 30 words) evaluating a week of teaching covering: {', '.join(valid_topics)} under syllabus {syllabus_code or 'standards'}. Focus on milestone achievements and skill growth."
        return client.models.generate_content(model='gemini-2.5-flash', contents=prompt).text.strip()
    except:
        return f"Weekly core subject objectives mapped to syllabus parameters completed successfully with steady skill advancement demonstrated across student tasks."


# --- Document Building Function ---
def build_word_document(week_no, start_date, end_date, days_data, conclusion):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)

    p = doc.add_paragraph()
    p.add_run("Weekly lesson evaluation for WEEK NO : ").font.bold = True
    p.add_run(f"   {week_no}  ").font.underline = True
    p.add_run("\t\t\t\t\t\tDATE: ").font.bold = True
    p.add_run(f" {start_date.strftime('%d/%m/%Y')} ").font.underline = True
    p.add_run(" to ")
    p.add_run(f" {end_date.strftime('%d/%m/%Y')} ").font.underline = True
    for r in p.runs: r.font.size = Pt(13)

    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    col_widths = [Inches(1.2), Inches(1.8), Inches(2.4), Inches(2.4), Inches(2.4)]
    headers = ["DATE/DAY", "TOPIC", "WHAT WENT WRONG", "EVEN BETTER WITH", "WAY FORWARD"]

    for i, title in enumerate(headers):
        table.rows[0].cells[i].text = title
        format_cell_text(table.rows[0].cells[i], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, day in enumerate(days_data):
        row_cells = table.rows[idx + 1].cells
        for col_idx, text in enumerate([day["date_str"], day["topic"], day["www"], day["ebw"], day["wf"]]):
            row_cells[col_idx].text = text

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell)
            format_cell_text(cell)

    doc.add_paragraph(
        "\nOverall evaluation is optional. For example, if you would like to record 'what went well' or other details of your lesson you may include at the end, as follows.").runs[
        0].font.size = Pt(9.5)

    bottom_table = doc.add_table(rows=2, cols=3)
    bottom_table.style = 'Table Grid'
    bottom_widths = [Inches(1.8), Inches(5.8), Inches(2.4)]

    bottom_table.rows[0].cells[0].text = "WHAT WENT WELL\n\n(Conclusion)"
    bottom_table.rows[0].cells[1].text = conclusion
    bottom_table.rows[1].cells[0].text = "REMARK"
    bottom_table.rows[1].cells[2].text = "\n\nName/Signature"

    set_cell_background(bottom_table.rows[0].cells[0], "66FFCC")
    set_cell_background(bottom_table.rows[1].cells[0], "66FFCC")

    for row in bottom_table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = bottom_widths[i]
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            align = WD_ALIGN_PARAGRAPH.RIGHT if i == 2 and row == bottom_table.rows[1] else None
            format_cell_text(cell, bold=(i == 0), alignment=align)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# --- Streamlit Web App Interface Execution ---
st.set_page_config(layout="wide", page_title="AI Lesson Evaluation Portal")
st.title("📋 A Week Lesson Integrated Evaluation Portal")

# --- MAIN AREA API KEY ENTRY ---
user_api_key = st.text_input(
    "🔑 Enter your Gemini API Key:", 
    type="password", 
    help="Get your personal API key from Google AI Studio using your Gmail account."
)

st.subheader("Configuration Panel")
c1, c2, c3 = st.columns(3)
week_input = c1.text_input("WEEK NO:", value="01")
start_dt = c2.date_input("Starting Monday Date:", datetime.date.today())
syllabus_input = c3.text_input("SYLLABUS CODE / SPECIFICATION:", value="9696")

day_offsets = {"Monday": 0, "Tuesday": 1, "Wednesday": 2, "Thursday": 3, "Saturday": 5}
end_dt = start_dt + datetime.timedelta(days=5)
st.info(
    f"Target Evaluation Window: **{start_dt.strftime('%d/%m/%Y')}** to **{end_dt.strftime('%d/%m/%Y')}** | Target Framework: **{syllabus_input if syllabus_input else 'General Curriculum'}**")

st.markdown("---")
st.write("##### Input Lesson Topics")
topics_collected = {}
col_l, col_r = st.columns(2)

for idx, (day_name, offset) in enumerate(day_offsets.items()):
    target_pane = col_l if idx < 3 else col_r
    calc_date = start_dt + datetime.timedelta(days=offset)
    date_str = calc_date.strftime("%d/%m/%Y")
    user_topic = target_pane.text_input(f"{day_name} ({date_str}) Topic:", key=f"inp_{day_name}")
    topics_collected[day_name] = {"date_str": f"{date_str}\n({day_name})", "topic": user_topic}

if st.button("GENERATE EVALUATION REPORT", type="primary"):
    if not user_api_key:
        st.error("❌ Key Configuration Error! Please input your personal Google Gemini API key at the top before generating.")
    else:
        try:
            # Initialize genai.Client dynamically with the user provided key
            client_instance = genai.Client(api_key=user_api_key)
            
            report_rows, raw_topics_list = [], []
            status_box = st.empty()

            processed_count = 0
            for day_name, d_info in topics_collected.items():
                topic_text = d_info["topic"]
                raw_topics_list.append(topic_text)

                if topic_text.strip():
                    if processed_count > 0:
                        for remaining in range(6, 0, -1):
                            status_box.info(f"⏳ Stabilizing system channels ({remaining}s)... Preparing: **{day_name}**")
                            time.sleep(1)

                    with st.spinner(f"🚀 Gemini analyzing: **{day_name}** ({topic_text})..."):
                        www_out, ebw_out, wf_out = generate_lesson_evaluation_with_gemini(client_instance, topic_text,
                                                                                          syllabus_input)
                    processed_count += 1
                else:
                    www_out, ebw_out, wf_out = "", "", ""

                report_rows.append(
                    {"date_str": d_info["date_str"], "topic": topic_text, "www": www_out, "ebw": ebw_out, "wf": wf_out})

            with st.spinner("✍️ Writing final summary box conclusion..."):
                final_conclusion = generate_weekly_conclusion_with_gemini(client_instance, raw_topics_list, syllabus_input)
            status_box.empty()

            st.session_state.update(
                {'data_processed': True, 'report_rows': report_rows, 'final_conclusion': final_conclusion})
                
        except Exception as api_err:
            st.error(f"❌ Connection or Authentication Error: {str(api_err)}")

if st.session_state.get('data_processed'):
    st.markdown("---")
    st.subheader("👁️ On-Screen Document Preview")
    st.table(st.session_state['report_rows'])
    st.write("**Generated Summary Box:**", st.session_state['final_conclusion'])

    word_file = build_word_document(week_input, start_dt, end_dt, st.session_state['report_rows'],
                                    st.session_state['final_conclusion'])
    st.download_button(label="📥 Download Official Word Document (.docx)", data=word_file,
                       file_name=f"Evaluation_Week_{week_input}.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
